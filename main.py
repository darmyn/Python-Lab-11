import os
import openpyxl
from docx import Document

PYTHON_LOGO_PATH = "assets/images/python-logo.png"
SPREADSHEET_PATH = "output/lab11sheet.xlsx"
DOCUMENT_PATH = "output/lab11report.docx"
TITLES = ["fish", "cheese", "car"]
PARAGRAPHS = [
    """
  Swimming gracefully through the crystal-clear waters of the 
  coral reef, the dazzling angelfish displays a mesmerizing 
  array of colors and patterns.
  """, """
  Aged cheddar, with its sharp and nutty flavor profile, pairs 
  exquisitely with the subtle sweetness of sliced apples on 
  a rustic cheese board.
  """, """
  Zooming along the scenic highway, the sleek sports car effortlessly hugs 
  the curves, delivering a thrilling driving experience.
  """
] 

if not os.path.exists(SPREADSHEET_PATH):
  workbook = openpyxl.Workbook()
  sheet = workbook.active
  for i in range(2):
    targetData = (i == 0) and PARAGRAPHS or TITLES
    for j in range(3):
      sheet.cell(row=j+1, column=i+1).value = targetData[j].capitalize()
  workbook.save(SPREADSHEET_PATH)

workbook = openpyxl.load_workbook(SPREADSHEET_PATH)
sheet = workbook.active

paragraphs_found = []
titles_found = []

for row in sheet.iter_rows():
  paragraph, title = row
  if paragraph:
    paragraphs_found.append(paragraph)
    titles_found.append(title)

paragraphs_found.sort(key=lambda x: len(x.value))
longest_paragraph = paragraphs_found[-1]
middle_paragraph = paragraphs_found[len(paragraphs_found) // 2]
shortest_paragraph = paragraphs_found[0]

doc = Document()

incorrectTitlesUsed = False
for title in titles_found:
  if not title.value.lower() in TITLES:
    incorrectTitlesUsed = True
    break

if incorrectTitlesUsed:
  print("Incorrect titles have been provided to the spreadsheet. Expecte titles to be: fish, cheese, car")
else:
  for heading, paragraph in zip(
      TITLES, [longest_paragraph, middle_paragraph, shortest_paragraph],
      strict=True):
  
    doc.add_heading(heading.capitalize(), level=1)
  
    p = doc.add_paragraph()
    for word in paragraph.value.split():
      if word.endswith("r") or word.endswith("r,") or word.endswith("r."):
        run = p.add_run(word)
        run.underline = True
      else:
        p.add_run(word)
      p.add_run(" ")
  
  doc.add_picture(PYTHON_LOGO_PATH)
  
  doc.save(DOCUMENT_PATH)
  